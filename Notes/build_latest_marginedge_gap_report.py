from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Notes/MEVS_vs_MarginEdge_Latest_Platform_Gap_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths, header_fill="E8EEF5"):
    set_table_width(table, widths)
    repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(11, 37, 69)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
    style_table(table, widths)
    doc.add_paragraph()
    return table


def add_h1(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def add_h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    run = p.add_run(label + ": ")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.text = "MEVS vs MarginEdge latest platform gap report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(85, 85, 85)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("MEVS vs MarginEdge Latest Platform Gap Report")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run("Remaining missing modules, unfinished workflows, and improvement priorities after the latest MEVS update.").italic = True

    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run("June 1, 2026  ")
    meta.add_run("Sources: ").bold = True
    meta.add_run("MarginEdge in-app browser inspection; MEVS local code, route, entity, and navigation review.")

    add_callout(
        doc,
        "Bottom line",
        "MEVS is much closer to MarginEdge than before. The newest update adds visible progress in accounting access, setup, transfers, receiving, count sheets, integrations, and performance. The biggest remaining risks are not broad module absence; they are incomplete workflows, broken or mismatched routing, placeholder screens, and missing persistence for setup operations.",
    )

    add_h1(doc, "Executive Summary")
    for item in [
        "Strongest progress: Accounting is now role-gated with org_owner, Orders exposes Transfers and Receiving, Inventory has count-sheet entities and UI, Integrations now persists records, and Performance reads POS sales data.",
        "Highest-priority blockers: Restaurant Setup exists in the navigation and as a page file, but it is not registered in pages.config or moduleConfig, so the module can be unreachable or entitlement-invisible.",
        "Most visible user-facing gaps: Accounting tabs, Performance tabs, Transfer/Receiving dialogs, Inventory count workflows, Labor scheduling, and Restaurant Setup screens still contain under-development placeholder content.",
        "Competitive opportunity: MEVS can stay ahead by turning AvT costing, AI Insights, integration health, and automated exception repair into first-class operator workflows rather than only matching MarginEdge module-for-module.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "MarginEdge Browser Baseline")
    doc.add_paragraph(
        "The current MarginEdge account shows a restaurant operations suite centered on Home, Orders, Performance, Vendors, Products, Recipes, SmartPrep, Inventory, Labor, Bill Pay, Card, Accounting, and Setup. Its competitive strength is not just module breadth; many modules are connected into daily operator workflows such as invoice approval, inventory count sheets, waste reporting, category performance, price movers, accounting mappings, and close-book controls."
    )

    add_table(
        doc,
        ["MarginEdge area", "Observed functions to match or beat"],
        [
            ("Home / Dashboard", "Sales, peer benchmark comparisons, period-to-date/year-to-date views, budget pacing, purchasing summary, price movers."),
            ("Orders", "Search orders, order placement, invoice approval, transfers, order setup."),
            ("Performance", "Budget, category report, controllable P&L, usage report, sales, forecast, price alerts, price movers, theoretical usage."),
            ("Inventory", "Summary, counts, waste log, waste summary, count sheets, products, liquor smart scale."),
            ("Accounting", "Export, reconciliation, sales entries, inventory entries, mappings for sales, vendors, PMIX, payment accounts, close books, budget setup."),
            ("Setup", "Users, shared devices, store groupings, integrations, POS, notifications."),
            ("Recipes / Prep", "Menu items, prepared items, bar items, menu analysis, recipe viewer, setup, SmartPrep."),
            ("Labor / Bill Pay / Card", "Labor summary, shifts, employees, setup; bill pay reconciliation/invoices/payments; purchase card capability."),
        ],
        [1.65, 4.85],
    )

    add_h1(doc, "What MEVS Improved in the Latest Update")
    add_table(
        doc,
        ["Area", "Latest MEVS progress", "Remaining concern"],
        [
            ("Accounting access", "moduleConfig now uses org_owner for Accounting, fixing the earlier org_admin role mismatch.", "Accounting sub-workflows remain mostly placeholders."),
            ("Setup", "RestaurantSetup.jsx was added with POS Setup, Store Groups, Shared Devices, Notifications, and Location Settings tabs.", "The page is not registered in pages.config and is absent from moduleConfig."),
            ("Orders", "Layout now exposes Transfers and Receiving under Orders; AutoOrdering has matching tabs.", "Dialogs still say workflow under development and appear not fully wired to Transfer/Receiving entities."),
            ("Inventory", "CountSheet and CountSession entities/migrations exist; UI now includes Count Sheets.", "Count-sheet creation and count session entry are placeholders; duplicate tab content needs cleanup."),
            ("Integrations", "Integration records now use persisted API entity operations instead of only mock UI state.", "Real OAuth, webhook status, sync health, and error recovery are still needed."),
            ("Performance", "Performance page now reads POS sales data and calculates key sales/budget cards.", "Budget is still synthetic and most analytical tabs are under development."),
        ],
        [1.55, 2.45, 2.5],
    )

    add_h1(doc, "Critical Fixes Before More Feature Work")
    add_table(
        doc,
        ["Priority", "Issue", "Evidence", "Recommended fix"],
        [
            ("P0", "Restaurant Setup routing and entitlement gap.", "Layout links to RestaurantSetup and RestaurantSetup.jsx exists, but pages.config and moduleConfig do not register it.", "Add lazy import and PAGES entry; add setup/restaurantSetup moduleConfig entry or attach it to Admin/Setup entitlement."),
            ("P0", "Accounting navigation tab mismatch.", "Layout links to Accounting?tab=mappings and Accounting?tab=close, while Accounting.jsx uses sales-mapping, vendor-mapping, pmix-mapping, and close-books.", "Change Layout to link to the exact Accounting tab values or update Accounting.jsx to accept aliases."),
            ("P0", "Inventory tab collision and unreachable waste summary.", "Inventory.jsx has two TabsContent blocks for count-sheets, and waste-summary content exists without a matching visible trigger.", "Merge the count-sheets content into one panel and add or remove the Waste Summary trigger intentionally."),
            ("P1", "Placeholder dialogs create false completeness.", "Transfer, Receiving, Count Sheet, Count Session, Restaurant Setup, Labor Setup, and Accounting mapping screens still display under-development messages.", "Convert placeholders into minimal end-to-end workflows before adding more tabs."),
        ],
        [0.6, 1.55, 2.0, 2.35],
    )

    add_h1(doc, "Remaining Missing or Incomplete Functions")
    add_table(
        doc,
        ["Module", "Still missing / incomplete vs MarginEdge", "Business impact"],
        [
            ("Performance", "Real controllable P&L, category report, price alerts, price movers, usage report, theoretical usage, reliable budget setup, and charts.", "Operators cannot yet diagnose margin leakage with the same depth as MarginEdge."),
            ("Orders", "Search Orders equivalent, real transfer creation, receiving against PO lines, partial receipt handling, discrepancy resolution, and order setup rules.", "Purchasing lifecycle is visible but not yet operationally complete."),
            ("Accounting", "Export, reconciliation, sales mapping, vendor mapping, PMIX mapping, sales entries, inventory entries, categories, payment accounts, and close-books locking.", "Finance users may see the module but cannot trust it as an accounting close workflow yet."),
            ("Inventory", "Count-sheet template builder, mobile count entry, count approvals, variance review, waste summary trigger, inventory products setup, and liquor smart scale equivalent.", "Inventory control exists in pieces but is not yet a daily count-and-variance system."),
            ("Labor", "Schedule builder, shift variance, POS/timeclock import, overtime rules, and labor setup persistence.", "Labor can summarize data but cannot yet manage labor operations end to end."),
            ("Setup", "Persistent POS setup, shared device registration, notification rules, location settings, and migrations for settings/devices/rules.", "Setup is currently more of a shell than a reliable configuration center."),
            ("Recipes / Prep", "Bar Items and SmartPrep-style production planning are still missing; recipe viewer/setup needs deeper workflow support.", "Menu and prep operations trail MarginEdge in bar and prep planning scenarios."),
            ("Bill Pay / Card", "MEVS has Bill Pay/Payments, but no MarginEdge Card equivalent was found.", "A purchase card workflow could become a differentiator if integrated into spend controls."),
        ],
        [1.3, 3.25, 1.95],
    )

    add_h1(doc, "Recommended Improvements to Stay Ahead")
    for item in [
        "AI action center: turn AI Insights into prioritized actions with owner, dollar impact, confidence, due date, and one-click task creation.",
        "Explainable AvT costing: show the variance source by invoice line, recipe usage, waste, sales mix, transfer, and count adjustment.",
        "Integration health console: expose sync status, failed records, retry actions, mapping drift, duplicate vendors/products, and POS/accounting connection health.",
        "Smart purchasing guardrails: recommend vendor substitutions, flag unusual price movement before orders are placed, and suggest order quantities based on forecast, par, and active prep plans.",
        "Close-readiness score: before accounting close, show missing invoices, unmapped categories, unreconciled payments, stale inventory counts, and unresolved receiving discrepancies.",
        "Setup completeness score: guide each restaurant through POS, vendors, products, recipes, accounts, users, devices, notifications, and first close with an auditable checklist.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "30-Day Implementation Roadmap")
    add_table(
        doc,
        ["Timeframe", "Focus", "Deliverables"],
        [
            ("Week 1", "Stabilize routing and navigation.", "Register RestaurantSetup; add moduleConfig support; fix Accounting tab links; clean Inventory duplicate count-sheets and Waste Summary tab."),
            ("Week 2", "Make Orders and Inventory workflows real.", "Create Transfer and Receiving records from dialogs; support partial receipts and discrepancies; create CountSheet templates; launch mobile count session entry."),
            ("Week 3", "Deepen Performance.", "Add real sales charts, budget setup, category report, price movers, usage report, theoretical usage, and P&L cards from existing entities."),
            ("Week 4", "Make Accounting and Setup credible.", "Build export/mapping/reconciliation minimum viable workflows; persist POS setup, shared devices, notification rules, store groups, and location settings."),
        ],
        [1.05, 2.0, 3.45],
    )

    add_h1(doc, "Implementation Checklist")
    for item in [
        "P0: Add RestaurantSetup to pages.config and moduleConfig.",
        "P0: Align Layout accounting links with Accounting.jsx tab values.",
        "P0: Merge duplicate count-sheets panels and expose or remove Waste Summary.",
        "P1: Wire Transfers and Receiving to api.entities.Transfer and api.entities.Receiving with create/list/update flows.",
        "P1: Add real CountSheet and CountSession creation, count entry, submit, and variance review.",
        "P1: Replace Accounting placeholder cards with minimum viable export, mapping, reconciliation, and close-book actions.",
        "P2: Implement Performance analytics with real charts, budget persistence, and alert thresholds.",
        "P2: Persist Restaurant Setup tabs with database tables for shared devices, notification rules, location settings, and POS configuration.",
        "P2: Add Bar Items and SmartPrep-style prep planning to Recipes/Menu Engineering.",
        "P3: Add purchase-card or controlled spend workflow to compete with MarginEdge Card.",
    ]:
        add_numbered(doc, item)

    add_h1(doc, "Conclusion")
    doc.add_paragraph(
        "The latest MEVS update moved the platform from broad feature coverage toward competitive parity, but several modules still stop at navigation-level or placeholder-level completeness. The immediate win is to fix the routing and tab issues, then convert the visible placeholder workflows into small but real end-to-end actions. Once that foundation is stable, MEVS can outpace MarginEdge by leaning into AI-driven exception handling, explainable AvT, and integration health features that reduce operator work instead of only reporting it."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
