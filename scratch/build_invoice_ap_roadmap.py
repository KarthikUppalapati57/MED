from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path


OUT = Path(r"C:\Users\ukart\OneDrive - University of Tennessee\M\INtern\MECURSOR\MEVS\Notes\RestOps_Invoice_AP_Implementation_Roadmap.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE7"
TEXT_GRAY = "4B5563"
GREEN = "E8F5EE"
GOLD = "FFF4D6"
RED = "FCE8E8"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table_geometry(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = 9360
    widths_dxa = [int(width * 1440) for width in widths_inches]
    widths_dxa[-1] += total - sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color=MID_GRAY, size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, 9, TEXT_GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    if "Callout" not in styles:
        style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Callout"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor.from_string(NAVY)
    style.paragraph_format.left_indent = Inches(0.16)
    style.paragraph_format.right_indent = Inches(0.16)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.10


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("IMPLEMENTATION ROADMAP")
    set_run(run, 10, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Invoice and Accounts Payable Operations")
    set_run(run, 25, NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Closing the operational gaps between RestOps and MarginEdge")
    set_run(run, 13, TEXT_GRAY, italic=True)

    metadata = [
        ("Prepared for", "RestOps Product and Engineering"),
        ("Prepared on", "June 14, 2026"),
        ("Document status", "Implementation planning baseline"),
        ("Primary objective", "Build a unified, production-ready invoice/AP lifecycle"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(metadata):
        set_cell_shading(table.cell(i, 0), LIGHT_GRAY)
        p1 = table.cell(i, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run(p1.add_run(label), 10, DARK_BLUE, bold=True)
        p2 = table.cell(i, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        set_run(p2.add_run(value), 10, NAVY)
    set_fixed_table_geometry(table, [1.55, 4.95])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(p, BLUE, "12", "5")


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), 11, NAVY, bold=True)
    set_run(p.add_run(text), 11, NAVY)
    set_fixed_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT_BLUE)
        p = header.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), 9.5, NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if status_col is not None and idx == status_col:
                val = str(value).lower()
                fill = GREEN if "strong" in val or "complete" in val else GOLD if "partial" in val or "weak" in val else RED
                set_cell_shading(cells[idx], fill)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == status_col:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), 9.5, NAVY, bold=(idx == 0))
    set_fixed_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_phase(doc, number, title, objective, workstreams, deliverable, dependencies=None):
    doc.add_heading(f"Phase {number}: {title}", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Objective. "), 11, NAVY, bold=True)
    p.add_run(objective)
    add_bullets(doc, workstreams)
    p = doc.add_paragraph()
    set_run(p.add_run("Phase deliverable. "), 11, NAVY, bold=True)
    p.add_run(deliverable)
    if dependencies:
        p = doc.add_paragraph()
        set_run(p.add_run("Key dependencies. "), 11, NAVY, bold=True)
        p.add_run(dependencies)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.38)
section.footer_distance = Inches(0.38)
configure_styles(doc)

header = section.header
hp = header.paragraphs[0]
hp.paragraph_format.space_after = Pt(2)
set_run(hp.add_run("RESTOPS  |  INVOICE & AP IMPLEMENTATION ROADMAP"), 8.5, TEXT_GRAY, bold=True)
set_paragraph_border_bottom(hp, MID_GRAY, "6", "2")
footer = section.footer
add_page_number(footer.paragraphs[0])

add_title_block(doc)

doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "RestOps already contains meaningful invoice, purchasing, receiving, inventory, and ledger capabilities. "
    "However, those capabilities are fragmented across the Invoices and Orders modules and do not yet provide "
    "the unified, accountant-ready operating experience visible in MarginEdge."
)
add_callout(
    doc,
    "Recommendation",
    "Make the Invoices module the unified AP system of record, connect it tightly to purchase orders, receiving, "
    "payment accounts, and ledger activity, and deliver the work in eight controlled phases.",
)

doc.add_heading("Current Position", level=2)
add_bullets(doc, [
    "RestOps is stronger in proactive purchasing automation: par-level replenishment, vendor price comparison, receiving, and inventory movement logic.",
    "MarginEdge is stronger in daily invoice/AP execution: invoice ledger, document review, reconciliation, payment-account visibility, approval, and bill-pay readiness.",
    "The largest RestOps weakness is not the absence of all functionality; it is fragmentation, shallow reconciliation, and limited operational visibility.",
])

doc.add_heading("Priority Outcomes", level=2)
add_numbered(doc, [
    "Create one unified invoice/AP ledger for the complete invoice lifecycle.",
    "Provide side-by-side invoice document review with trustworthy ingestion and duplicate controls.",
    "Implement durable vendor/item mappings and line-level three-way reconciliation.",
    "Introduce policy-driven approvals, payment accounts, and a clear bill-pay lifecycle.",
    "Add exports, reporting, auditability, and operational monitoring.",
])

doc.add_page_break()
doc.add_heading("1. Gap Assessment", level=1)
doc.add_paragraph(
    "The following assessment distinguishes between capabilities that are absent, partially implemented, or materially weaker than the target MarginEdge-style operating model."
)
gap_rows = [
    ("Invoice ledger", "Partial", "Invoice table exists, but lacks AP status buckets, aging, action reasons, payment-account state, and batch actions."),
    ("Upload intake", "Partial", "OCR upload exists; ingestion history, retry controls, multi-document review, and duplicate resolution need depth."),
    ("Email intake", "Partial / risky", "IMAP configuration exists, but credential handling and ingestion monitoring require production hardening."),
    ("Invoice approval", "Partial", "Approve/reject exists; configurable chains, escalation, comments, delegation, and batch approval are missing."),
    ("Bill-pay linkage", "Partial", "Approval can create ledger bills, but the user-facing invoice-to-payment lifecycle is incomplete."),
    ("Payment accounts", "Missing", "No invoice-level selection of AP, bank, credit card, petty cash, transfer, or custom accounts."),
    ("Invoice photo review", "Weak", "Files can be accessed, but there is no integrated multi-page viewer with zoom, rotate, or field verification."),
    ("CSV export", "Missing", "No operational invoice ledger, line-item, variance, or payment CSV exports."),
    ("Category summary", "Missing", "No invoice-level category/GL summary or split-coding review experience."),
    ("Line-item reconciliation", "Weak", "Matching is primarily total-level; true PO-receipt-invoice line reconciliation is absent."),
    ("Vendor/item matching", "Weak", "Matching relies heavily on names; durable aliases, SKUs, packaging, and confirmed mappings are absent."),
]
add_table(doc, ["Capability", "RestOps status", "Primary weakness"], gap_rows, [1.35, 1.15, 4.0], status_col=1)

doc.add_heading("Structural Weaknesses", level=2)
add_bullets(doc, [
    "Invoices and order-based approval are split across separate user experiences.",
    "Invoice matching lacks a durable vendor-item identity layer.",
    "Operational statuses do not clearly communicate why an invoice requires action.",
    "Approval and payment steps are implemented as actions, not as a governed workflow.",
    "Accounting review is line-item heavy because category summary and split coding are missing.",
    "The platform lacks a complete audit trail that connects source document, PO, receipt, approval, bill, and payment.",
])

doc.add_page_break()
doc.add_heading("2. Target Operating Model", level=1)
add_callout(
    doc,
    "Target state",
    "Every invoice enters one processing queue, is validated and matched, resolves exceptions, passes the correct approval policy, creates an AP bill, and closes through a traceable payment event.",
)
doc.add_heading("End-to-End Workflow", level=2)
workflow_rows = [
    ("1", "Intake", "Upload, email, EDI, recurring invoice, or manual entry", "Processing"),
    ("2", "Extract and validate", "OCR, field validation, duplicate detection, source preservation", "Processing / Action Required"),
    ("3", "Match", "Vendor, item, PO, receipt, price, quantity, packaging", "Matched / Variance"),
    ("4", "Resolve", "User fixes missing mappings, mismatches, and coding exceptions", "Action Required"),
    ("5", "Approve", "Policy-driven approval, comments, delegation, and escalation", "Pending Approval / Approved"),
    ("6", "Create bill", "Create or update AP bill and accounting coding", "Approved / Ready for Payment"),
    ("7", "Pay", "Choose payment account, schedule, record, or reconcile payment", "Scheduled / Paid"),
    ("8", "Close", "Finalize audit trail, export, and reporting", "Closed"),
]
add_table(doc, ["Step", "Stage", "System behavior", "Primary status"], workflow_rows, [0.55, 1.25, 3.7, 1.0])

doc.add_heading("Unified Ledger Views", level=2)
add_bullets(doc, [
    "Processing: invoices being extracted, validated, or matched.",
    "Action Required: duplicates, missing vendor/item mapping, missing PO/receipt, or reconciliation variance.",
    "Pending Approval: invoices waiting for one or more required approvers.",
    "Approved: approved invoices with AP bills created.",
    "Scheduled for Payment: bills assigned to payment accounts and dates.",
    "Paid and Closed: completed invoices retained for audit, reporting, and export.",
])

doc.add_heading("Required Invoice Detail Tabs", level=2)
add_table(doc, ["Tab", "Purpose"], [
    ("Documents", "Original invoice photos/PDFs, thumbnails, source metadata, and review tools."),
    ("Line Items", "Extracted invoice lines, product mappings, quantities, units, packaging, and prices."),
    ("Category Summary", "Category, GL account, allocation, taxes, fees, credits, and split coding."),
    ("Reconciliation", "PO-receipt-invoice comparison and explainable variance resolution."),
    ("Approval History", "Policy, approvers, comments, decisions, reminders, and escalation."),
    ("Payment History", "AP bill, payment account, schedule, payment events, failures, and reconciliation."),
    ("Audit Trail", "Complete record of user, timestamp, old value, new value, reason, and source."),
], [1.55, 4.95])

doc.add_page_break()
doc.add_heading("3. Phased Implementation Plan", level=1)
doc.add_paragraph(
    "Delivery should prioritize operational coherence before advanced automation. Each phase must be usable independently while establishing the foundation for the next phase."
)

add_phase(doc, 1, "Unified Invoice/AP Ledger",
    "Create one operational workspace that becomes the system of record for the complete invoice lifecycle.",
    [
        "Consolidate invoice approval and AP status management into the Invoices module.",
        "Add status buckets: Processing, Action Required, Pending Approval, Approved, Scheduled, Paid, Closed, and Rejected.",
        "Add filters for dates, vendor, PO, match status, approval status, payment status, payment account, location, and aging.",
        "Add batch actions for approval, reviewer assignment, payment-account assignment, export, close, and bill-pay routing.",
        "Add action-required reason codes and clear next-action guidance.",
    ],
    "An accountant-ready invoice ledger that provides complete operational visibility.",
    "Canonical invoice statuses, permissions, tenant scoping, and migration of existing invoice/order approval behavior."
)

add_phase(doc, 2, "Document Intake and Photo Review",
    "Make invoice intake reliable and allow reviewers to validate invoices without leaving the application.",
    [
        "Support drag-and-drop PDFs/images, mobile photos, multi-page documents, and multiple attachments.",
        "Add dedicated invoice email intake, attachment processing, ingestion job history, retries, and source tracking.",
        "Implement duplicate detection using vendor, invoice number, date, amount, and document fingerprint.",
        "Build a side-by-side document viewer with zoom, rotate, page navigation, thumbnails, and OCR confidence indicators.",
        "Preserve original documents, extraction versions, and reviewer corrections.",
        "Replace raw IMAP password storage with encrypted secrets or provider OAuth.",
    ],
    "A secure, observable intake pipeline and a production-ready invoice review workspace.",
    "Object storage, secret management, processing-job monitoring, OCR metadata, and mobile upload behavior."
)

add_phase(doc, 3, "Vendor and Item Matching",
    "Create durable product identity and continuously improve future invoice processing.",
    [
        "Add vendor aliases, vendor account identifiers, vendor items, product packaging, and mapping history.",
        "Match by vendor account number, alias, vendor SKU, UPC/external ID, packaging/unit, then normalized description.",
        "Display match suggestions and confidence scores.",
        "Allow users to confirm a match, search products, or create a new product.",
        "Persist confirmed mappings and flag packaging, unit, and price changes.",
    ],
    "A reusable vendor-item mapping layer that reduces manual review over time.",
    "Vendor master quality, product master quality, normalization rules, and mapping permissions."
)

add_phase(doc, 4, "Three-Way Line-Item Reconciliation",
    "Reconcile purchase order, physical receipt, and invoice at line-item level.",
    [
        "Calculate ordered, received, and invoiced quantities for every matched item.",
        "Compare PO price, invoice price, quantity, packaging, and extended totals.",
        "Assign statuses: exact match, within tolerance, price variance, quantity variance, missing receipt, missing PO, unexpected item, duplicate item, and packaging mismatch.",
        "Support tolerances by organization, vendor, category, and item.",
        "Require resolution notes or approval overrides for exceptions.",
    ],
    "Explainable line-level reconciliation that supports safe automation and faster approvals.",
    "Stable PO and receiving line identities, vendor-item mappings, and tolerance configuration."
)

doc.add_page_break()
add_phase(doc, 5, "Category Summary and Accounting Review",
    "Enable efficient invoice coding and review without editing every individual line.",
    [
        "Add category and GL summaries for subtotal, taxes, delivery/fuel charges, credits, and other charges.",
        "Support inventory category, accounting category, GL account, and location allocation.",
        "Allow split coding across categories, GL accounts, and locations.",
        "Expose uncategorized totals and coding exceptions.",
        "Propagate approved coding to ledger bills and accounting exports.",
    ],
    "A concise accounting review surface with complete coding visibility.",
    "Chart of accounts, category mappings, allocation rules, and accounting permissions."
)

add_phase(doc, 6, "Configurable Approval Workflow",
    "Replace simple approve/reject actions with governed, policy-driven approvals.",
    [
        "Create approval rules based on total, vendor, category, location, variance, missing PO/receipt, and payment account.",
        "Support single-step, sequential, and parallel approvals.",
        "Add batch approval, comments, request changes, delegation, reminders, and escalation.",
        "Record the policy evaluation and every approval event in the audit trail.",
        "Auto-approve invoices only when matching and policy requirements are satisfied.",
    ],
    "Configurable approval chains with clear accountability and escalation.",
    "Role/permission model, notification services, workflow event logging, and audit history."
)

doc.add_heading("Example Approval Policy", level=3)
add_table(doc, ["Condition", "Required action"], [
    ("Under $250 and fully matched", "Auto-approve"),
    ("$250 to $2,500", "Location manager approval"),
    ("Over $2,500", "Location manager plus organization owner"),
    ("Variance above 5%", "Accounting review"),
    ("Missing PO or receipt", "Action Required; cannot auto-approve"),
], [2.5, 4.0])

add_phase(doc, 7, "Payment Accounts and Bill Pay",
    "Provide a clear, traceable invoice-to-payment lifecycle.",
    [
        "Add payment-account management for AP, checking, credit card, petty cash, internal transfer, vendor auto-pay, and custom accounts.",
        "Show bill-pay eligibility, payment account, method, due date, scheduled date, partial payment, reference, status, and failure reason.",
        "Create or update an AP bill on invoice approval.",
        "Support scheduling, marking externally paid, partial payments, and payment reconciliation.",
        "Create ledger entries and close invoices only after valid payment events.",
    ],
    "A complete bill-pay workflow connected to invoices, AP bills, payment accounts, and ledger entries.",
    "Payment provider strategy, banking/security controls, ledger model, and reconciliation rules."
)

add_phase(doc, 8, "Exports, Reporting, and Audit",
    "Make AP operations measurable, portable, and defensible.",
    [
        "Add CSV exports for ledger, line items, category summary, approval history, reconciliation variance, payments, and vendor price history.",
        "Add AP aging, upcoming payments, vendor spend, variance, unmatched invoice, approval-cycle, and ingestion-failure reports.",
        "Record user, timestamp, old/new values, reason, source, and all related entities for every material action.",
        "Add operational dashboards for processing failures, queue aging, and unresolved exceptions.",
    ],
    "Complete reporting, export, and audit coverage for operations and accounting.",
    "Unified workflow events, consistent identifiers, data retention policy, and export permissions."
)

doc.add_page_break()
doc.add_heading("4. Proposed Data Model Additions", level=1)
doc.add_paragraph(
    "The existing invoices, auto_orders, receivings, transfers, ledger bills, and payments should remain, but the following entities are needed to support production-grade AP operations."
)
data_rows = [
    ("invoice_documents", "Original documents, pages, thumbnails, storage paths, hashes, source, and versions."),
    ("invoice_ingestion_jobs", "Email/upload processing state, attempts, errors, timing, and retry metadata."),
    ("invoice_action_reasons", "Structured reasons that place an invoice in Action Required."),
    ("vendor_aliases", "Normalized names and identifiers that resolve to canonical vendors."),
    ("vendor_items", "Vendor-specific SKU, description, packaging, unit, and latest pricing."),
    ("vendor_item_mappings", "Durable vendor-item-to-product relationships with confidence and confirmation history."),
    ("invoice_line_matches", "Invoice line, PO line, receiving line, mapping confidence, and match status."),
    ("reconciliation_variances", "Price, quantity, packaging, total variance, tolerance, resolution, and approver."),
    ("approval_policies", "Policy conditions, scope, priority, and activation state."),
    ("approval_instances", "Evaluated workflow for a specific invoice."),
    ("approval_steps", "Approvers, order, decision, comment, delegation, reminder, and escalation."),
    ("payment_accounts", "Account type, provider reference, display details, eligibility, and tenant scope."),
    ("invoice_allocations", "Category, GL account, location, amount, percentage, and source."),
    ("invoice_audit_events", "Immutable invoice-related workflow and data-change history."),
]
add_table(doc, ["Entity", "Purpose"], data_rows, [2.05, 4.45])

doc.add_heading("Security and Integrity Requirements", level=2)
add_bullets(doc, [
    "Apply organization, brand, and location scoping consistently through RLS and service-layer checks.",
    "Encrypt email and payment integration secrets; never store raw credentials in general metadata.",
    "Use immutable audit events for approvals, coding overrides, payment changes, and reconciliation resolutions.",
    "Require idempotency for ingestion, bill creation, payment creation, and ledger posting.",
    "Prevent invoice closure when required approval, payment, or reconciliation conditions remain incomplete.",
])

doc.add_heading("5. UX and Navigation Recommendation", level=1)
add_callout(
    doc,
    "Product decision",
    "Use Invoices as the unified AP workspace. Keep Orders focused on purchase orders, receiving, and transfers, while linking directly into the corresponding invoice/AP records.",
)
add_table(doc, ["Area", "Recommended responsibility"], [
    ("Invoices / AP", "Invoice ledger, intake, document review, matching, reconciliation, approvals, coding, bills, payment status, exports, and audit."),
    ("Orders", "Purchase-order generation, approval, vendor sending, receiving, and internal transfers."),
    ("Payments", "Payment execution, account management, payment history, failures, and reconciliation."),
    ("Accounting", "GL mappings, exports, close controls, and ledger-level reconciliation."),
], [1.45, 5.05])

doc.add_page_break()
doc.add_heading("6. Delivery Roadmap", level=1)
roadmap_rows = [
    ("Release 1", "Weeks 1-3", "Unified AP ledger, canonical statuses, filters, action reasons, batch foundations", "Operational visibility"),
    ("Release 2", "Weeks 4-6", "Document viewer, multi-page intake, email hardening, ingestion history, duplicates", "Reliable intake and review"),
    ("Release 3", "Weeks 7-10", "Vendor/item mappings and three-way line reconciliation", "Core differentiation"),
    ("Release 4", "Weeks 11-13", "Category summary, split coding, configurable approvals", "Accounting and governance"),
    ("Release 5", "Weeks 14-16", "Payment accounts, bill-pay lifecycle, exports, reporting, audit refinement", "End-to-end AP completion"),
]
add_table(doc, ["Release", "Timing", "Scope", "Outcome"], roadmap_rows, [0.85, 0.9, 3.7, 1.05])

doc.add_heading("Recommended Team Shape", level=2)
add_bullets(doc, [
    "Product lead: owns AP operating model, workflow policy, and release acceptance.",
    "Frontend engineer: ledger, document review, reconciliation, coding, and approval experiences.",
    "Backend/data engineer: schema, RLS, ingestion, matching, reconciliation, workflow, and ledger integrity.",
    "Integration engineer: email intake, object storage, payment providers, and accounting exports.",
    "QA/operations partner: fixture design, role testing, accounting validation, and production-readiness checks.",
])

doc.add_heading("Implementation Principles", level=2)
add_numbered(doc, [
    "Build one canonical workflow rather than duplicating invoice behavior across modules.",
    "Make every exception explainable and actionable.",
    "Automate only when matching confidence and policy allow it.",
    "Preserve the original source document and every material decision.",
    "Use idempotent workflow services for bills, payments, ledger entries, and ingestion.",
    "Deliver complete vertical slices and test them with realistic invoice, PO, receipt, and payment fixtures.",
])

doc.add_heading("7. Risks and Mitigations", level=1)
risk_rows = [
    ("Fragmented ownership", "Features continue to diverge across Orders, Invoices, Payments, and Accounting.", "Define Invoices/AP as system of record and publish module contracts."),
    ("Poor master data", "Vendor/item matching confidence remains low.", "Introduce aliases, confirmed mappings, packaging identity, and correction history."),
    ("Unsafe automation", "Incorrect invoices are auto-approved or paid.", "Gate automation on reconciliation, policy, tolerances, and audit checks."),
    ("Duplicate financial events", "Repeated jobs create duplicate bills, payments, or ledger entries.", "Use idempotency keys and unique source references."),
    ("Credential exposure", "Email/payment secrets are stored insecurely.", "Use encrypted secrets, OAuth, provider tokens, and limited service access."),
    ("Migration inconsistency", "Existing invoice/order statuses do not map cleanly.", "Create canonical status mapping, backfill scripts, and migration reports."),
]
add_table(doc, ["Risk", "Impact", "Mitigation"], risk_rows, [1.4, 2.35, 2.75])

doc.add_page_break()
doc.add_heading("8. Definition of Done", level=1)
doc.add_paragraph(
    "The program should be considered complete only when the workflow is operationally coherent, financially safe, and testable end to end."
)
acceptance = [
    "An invoice from upload or email appears in the unified ledger with source, processing state, original document, and ingestion history.",
    "A reviewer can compare the document with extracted fields and line items without downloading it.",
    "Vendor and item matches persist and improve future invoice processing.",
    "Every matched invoice line displays ordered, received, invoiced, and variance values.",
    "Action Required clearly identifies the problem and the exact resolution path.",
    "Approval policies select the correct approvers and retain a complete decision history.",
    "Approved invoices create idempotent AP bills with category/GL allocations.",
    "Payment accounts and payment states are visible from invoice detail.",
    "Valid payments create ledger entries and transition invoices to paid/closed.",
    "CSV exports, reports, and audit records are accurate and permission-scoped.",
    "Role, RLS, migration, retry, duplicate, and failure-path tests pass with realistic fixtures.",
]
for item in acceptance:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_heading("9. Immediate Next Actions", level=1)
add_numbered(doc, [
    "Approve the unified AP operating model and make Invoices the canonical AP workspace.",
    "Finalize canonical statuses, action-required reason codes, and the invoice lifecycle state machine.",
    "Create schema migrations for document intake, vendor/item matching, reconciliation, approval policies, payment accounts, and audit events.",
    "Build the unified ledger and side-by-side invoice review experience as the first production release.",
    "Create realistic end-to-end fixtures covering upload, email, duplicate, PO match, receipt variance, approval, bill creation, and payment.",
])

add_callout(
    doc,
    "Highest-value first release",
    "Unified AP ledger + side-by-side document review + durable vendor/item matching + line-level three-way reconciliation.",
    GREEN,
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
