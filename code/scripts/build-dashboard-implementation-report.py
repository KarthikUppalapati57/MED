from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "role_based_dashboard_implementation_report.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "4B5563"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D1D5DB", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Restops 360 Role-Based Dashboard Implementation Report")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(INK)
    r.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Home-style restaurant operator dashboard, role scoping, KPIs, reporting, scheduler, and QA status")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level <= 2 else DARK_BLUE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [Inches(6.5)])
    set_table_borders(table, color="DADCE0", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        set_cell_text(cell, header, bold=True, color=INK, size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8.5)
    doc.add_paragraph()
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in [
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Restops 360 dashboard implementation")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


phase_rows = [
    ("1", "Discovery and benchmark", "Compared MarginEdge Home with Restops Dashboard and defined the restaurant operator dashboard direction.", "Complete"),
    ("2", "Role-based dashboard plan", "Mapped org owner, brand manager, location manager, ground staff, and platform admin dashboard expectations.", "Complete"),
    ("3", "Supabase dashboard data foundation", "Added SQL/RPC support for role-scoped dashboard summaries and confirmed Supabase changes were applied.", "Complete"),
    ("4", "Org owner operator dashboard", "Built organization dashboard with platform workflow visibility plus restaurant performance KPIs.", "Complete"),
    ("5", "Brand manager operator dashboard", "Added brand-scoped performance dashboard for brand managers, using brand context and brand KPIs.", "Complete"),
    ("6", "Location manager daily dashboard", "Added location-scoped daily dashboard for sales, COGS, labor, AP, inventory, and action items.", "Complete"),
    ("7", "Ground staff dashboard", "Added staff dashboard limited to assigned modules and role permissions.", "Complete"),
    ("8", "Data health and onboarding signals", "Added Data Health Score, source coverage, and onboarding prompts for missing POS, AP, inventory, labor, and product data.", "Complete"),
    ("9", "Manager Decision Brief", "Added plain-language business readout with primary risk, handoff note, and action summary.", "Complete"),
    ("10", "Forecast Intelligence", "Added projected sales, prime cost pressure, inventory risk, sales trend, weekday factor, and volatility signals.", "Complete"),
    ("11", "Executive Report", "Added owner-ready scorecard combining performance, forecast, workflow, unresolved actions, and review notes.", "Complete"),
    ("12", "Dashboard persistence", "Persisted action status, handoff notes, review logs, report preferences, rules, and delivery records.", "Complete"),
    ("13", "Scheduled Reports", "Added daily and weekly dashboard report preferences with recipient roles and content toggles.", "Complete"),
    ("14", "Report Delivery Log", "Added delivery log showing report status, recipients, notifications, and retry metadata.", "Complete"),
    ("15", "Production Readiness panel", "Added checks for data sources, scheduler state, recipient coverage, notification health, and rules state.", "Complete"),
    ("16", "Dashboard Rules Center", "Added role/scope rules, COGS/labor/prime-cost targets, escalation thresholds, and notification behavior.", "Complete"),
    ("17", "Manual report actions", "Added Send Daily Now and Send Weekly Now actions for manager roles.", "Complete"),
    ("18", "Scheduler Edge Function", "Built and deployed Supabase function dashboard-report-scheduler for report generation and notification delivery.", "Complete"),
    ("19", "Audit and notifications", "Wired report events to notification and audit flows so scheduled/manual actions are visible to users.", "Complete"),
    ("20", "Permission hardening", "Confirmed manager controls are enabled only for org, brand, and location managers; staff stays read-only for dashboard operations.", "Complete"),
    ("21", "Production verification", "Verified deployed scheduler produced daily and weekly sent records with notifications and cleared stale success metadata.", "Complete"),
    ("22", "QA tenant/account seed", "Created repeatable seed data for multiple organizations, brands, locations, users, roles, and report preferences.", "Complete"),
    ("23", "Role QA on production", "Logged in with QA accounts and verified platform admin, org owner, brand manager, location manager, and staff dashboard behavior.", "Complete"),
]


role_rows = [
    ("Platform Admin", "Global platform", "Platform Overview", "Organizations, users, subscriptions, revenue, audit activity", "Platform admin pages only"),
    ("Organization Owner", "Organization", "QA Bistro Group Dashboard", "Restaurant KPIs, forecasts, reports, rules, readiness, action center", "Manager controls enabled"),
    ("Brand Manager", "Brand", "North Fork Grill Dashboard", "Brand WTD sales, prime cost, open orders, low stock, forecasts", "Manager controls enabled"),
    ("Location Manager", "Location", "North Fork Downtown Dashboard", "Today's sales, COGS, labor, action items, readiness", "Manager controls enabled"),
    ("Ground Staff", "Assigned location/modules", "My Dashboard", "My uploads, pending invoices, assigned modules, shift plan", "No report/rules/manager controls"),
]


kpi_rows = [
    ("Business performance", "Period sales, WTD sales, today's sales, gross margin, COGS, labor, prime cost"),
    ("Workflow pressure", "Unpaid AP, pending invoices, open orders, low stock, action items, exceptions"),
    ("Forecasting", "Week sales forecast, month sales forecast, prime cost forecast, inventory risk, sales trend, weekday factor, volatility"),
    ("Operational readiness", "Data Health Score, connected data sources, POS setup, budget targets, product coverage, report readiness"),
    ("Reporting", "Scheduled daily/weekly reports, recipients, delivery status, notification count, resend metadata"),
    ("Staff execution", "Assigned module count, module task queues, shift checklist, today's urgent alerts"),
]


technical_rows = [
    ("Frontend", "src/pages/Dashboard.jsx", "Role-specific dashboard variants, KPI panels, forecast/report/rules/readiness components."),
    ("Layout/navigation", "src/Layout.jsx", "Role-aware and permission-aware navigation filtering."),
    ("Permissions", "src/hooks/usePermissions.jsx", "Role hierarchy and dashboard access checks."),
    ("Notifications", "src/lib/notificationService.js", "Report delivery notification support."),
    ("Audit", "src/lib/audit.js", "Dashboard report and scheduler audit events."),
    ("Database", "supabase/migrations/075-077", "Dashboard summary RPC, persistence tables, scheduler/report delivery schema."),
    ("Edge Function", "supabase/functions/dashboard-report-scheduler/index.ts", "Daily/weekly report scheduler and delivery writer."),
    ("QA seed", "scripts/seed-role-qa-data.mjs", "Multi-organization role QA seed users, memberships, permissions, and report preferences."),
]


qa_rows = [
    ("qa.platform.admin@restops.test", "Platform Admin", "Verified Platform Overview and platform admin navigation."),
    ("qa.owner.bistro@restops.test", "Org Owner", "Verified org dashboard with operator KPIs, reports, rules, readiness, and manager controls."),
    ("qa.brand.northfork@restops.test", "Brand Manager", "Verified North Fork brand dashboard with brand scope and manager controls."),
    ("qa.location.northfork@restops.test", "Location Manager", "Verified North Fork Downtown daily dashboard with location scope and manager controls."),
    ("qa.staff.northfork@restops.test", "Ground Staff", "Verified staff dashboard with assigned modules and no manager/report/rules controls."),
]


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "Implementation status",
        "The role-based restaurant operator dashboard is implemented and production-verified on restops-360.com. "
        "The current system gives each role the correct dashboard depth: platform admins see SaaS operations, managers see platform workflow plus restaurant performance, and ground staff see only assigned module work.",
    )

    add_heading(doc, "Executive Summary")
    add_body(
        doc,
        "We transformed the Restops Dashboard into a Home-style operating command center inspired by MarginEdge, but adapted for Restops 360's multi-tenant role model. "
        "The dashboard now combines platform workflow visibility, restaurant performance KPIs, forecasts, operational readiness, rules, scheduled reporting, and role-specific action surfaces.",
    )
    add_body(
        doc,
        "The implementation is role scoped. Organization owners see organization-level operations and performance, brand managers see their brand, location managers see their location, and ground staff see only the modules they can access. Platform admins retain a SaaS platform operations dashboard.",
    )

    add_heading(doc, "Phase Implementation Record")
    add_matrix(
        doc,
        ["Phase", "Name", "What was implemented", "Status"],
        phase_rows,
        [Inches(0.55), Inches(1.55), Inches(3.85), Inches(0.55)],
    )

    add_heading(doc, "Role-Based Dashboard Behavior")
    add_matrix(
        doc,
        ["Role", "Scope", "Dashboard", "Main KPI / workflow surface", "Controls"],
        role_rows,
        [Inches(1.2), Inches(1.0), Inches(1.35), Inches(2.25), Inches(0.7)],
    )

    add_heading(doc, "KPI and Operating Signal Catalog")
    add_matrix(
        doc,
        ["Area", "Signals"],
        kpi_rows,
        [Inches(1.75), Inches(4.75)],
    )

    add_heading(doc, "Major Product Capabilities Added")
    for item in [
        "Home-style operator dashboard with role-specific KPI language and scope labels.",
        "Data Health Score and onboarding prompts for incomplete operational data.",
        "Manager Decision Brief for quick leadership handoff.",
        "Forecast Intelligence for sales, prime cost, inventory risk, trend, weekday factor, and volatility.",
        "Executive Report with scorecard, unresolved actions, and manager review notes.",
        "Dashboard Rules Center for thresholds, escalation behavior, and recipient role settings.",
        "Scheduled Reports for daily and weekly reporting, with manual send actions.",
        "Report Delivery Log with sent/failed/skipped state and notification counts.",
        "Production Readiness panel to validate data, scheduler, reports, recipients, and rules.",
        "Ground staff dashboard filtered by module permissions.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Technical Assets")
    add_matrix(
        doc,
        ["Area", "Asset", "Purpose"],
        technical_rows,
        [Inches(1.1), Inches(2.1), Inches(3.3)],
    )

    add_heading(doc, "Production QA Result")
    add_body(
        doc,
        "Production testing was performed on restops-360.com using seeded QA accounts across multiple organizations, brands, and locations. "
        "High-privilege QA accounts required MFA, which confirmed the production security gate is active.",
    )
    add_matrix(
        doc,
        ["Account", "Role", "Verified result"],
        qa_rows,
        [Inches(2.2), Inches(1.25), Inches(3.05)],
    )

    add_heading(doc, "Known Follow-Up Recommendations")
    for item in [
        "Keep the scheduler Edge Function deployed after every scheduler code change.",
        "Keep QA seed data available for regression testing role scope after permission or dashboard changes.",
        "Configure production Sentry DSN and PostHog key if product telemetry is required.",
        "Add richer demo data for sales, labor, invoices, inventory, and orders so screenshots show realistic KPI movement.",
        "Add automated end-to-end tests for the five role dashboards to prevent scope regressions.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Current Conclusion")
    add_body(
        doc,
        "The dashboard implementation is not deviating from the original plan. It follows the planned role-based Home/dashboard model: managers receive both platform operations and restaurant performance, while staff receive module-specific execution views. "
        "The system is ready for realistic demo data, automated regression tests, and ongoing KPI refinement.",
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
